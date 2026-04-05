from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_task(num, name, code, approved, is_new=False):
    tag = ' — NEW' if is_new else ''
    p = doc.add_paragraph()
    run = p.add_run(f'Task #{num} — {name} ({code}){tag}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p2 = doc.add_paragraph(f'Approved: {"Yes" if approved else "No"}')
    p2.paragraph_format.space_after = Pt(2)

def field(text):
    doc.add_paragraph(text, style='List Bullet')

def fields_label():
    p = doc.add_paragraph('Fields:')
    p.runs[0].bold = True

def doc_header(name):
    p = doc.add_paragraph()
    run = p.add_run(f'Document: {name}')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 110, 195)
    doc.add_paragraph('Document Fields:')

def no_docs():
    doc.add_paragraph('Documents: None')

def sep():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

# TITLE
h = doc.add_heading('SHIPSY EXIM - MANAGE TASKS\nCOMPLETE TASK LIST', level=0)
h.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph('Total: 40 Tasks | 4 Documents (Tasks #19, #24, #25, #26 — all FF)')
doc.add_paragraph('')

# ═══ SHIPPER ═══
doc.add_heading('SHIPPER (Ops) — 17 Tasks', level=2)
doc.add_heading('DRAFTS', level=3)

add_task(1, 'Select Mode Of Shipment', 'SMT25', False)
fields_label()
field('Mode (SMF87) — dropdown, required — Options: FCL, LCL, AIR, BULK (MR), BREAK BULK (MB)')
field('Incoterm (SMF88) — dropdown, required — Options: EXW, FCA, CPT, CIP, DAP, DPU, DDP, FAS, FOB, CFR, CIF')
field('Spot / Normal — dropdown, required, NEW — Options: Spot, Normal')
no_docs(); sep()

add_task(2, 'Select Port Details', 'SMT7', False)
fields_label()
field('Place of Receipt at Origin (SMF103) — dropdown, required')
field('Port of Loading (SMF4) — dropdown, required')
field('Port of Discharge (SMF5) — dropdown, required')
field('Destination Port (SMF3) — dropdown, required')
no_docs(); sep()

add_task(3, 'Enter Container Details', 'SMT8', False)
fields_label()
field('Container Details (SMF9) — text, required')
field('Container Type (SMF10) — dropdown')
field('Container Size (SMF11) — dropdown')
field('Container Count (SMF12) — number')
field('Container Total Weight (SMF13) — number')
field('Container Weight Unit (SMF14) — dropdown')
no_docs(); sep()

add_task(4, 'Vendor Selection', 'TBD', False, True)
fields_label()
field('Vendor Selection — addmore, required — Options: Freight Forwarder, Shipping Line, CHA, CFS, ICD, Break Bulk Vendor, Surveyor, Transporter')
field('Segment — dropdown — Options: DPD, DPD CFS, Non DPD (If CFS/ICD selected)')
field('Category — dropdown — Options: General In-gauge, General Out-gauge, Haz cargo, Reefer (If CFS/ICD selected)')
field('Destuff Indicator — dropdown — Options: Loaded, Destuffed (If CFS/ICD selected)')
field('Panel Identifier — dropdown — Options: Panel Lines (Off dock), Panel Lines (Non off dock), Non Panel Lines (If CFS/ICD selected)')
field('Buffer Days — dropdown — Options: 7 days, 14 days, 21 days (If CFS/ICD selected)')
no_docs(); sep()

add_task(5, 'Run Global Plan Optimizer', 'SMT10', True)
fields_label()
field('Bid Details (SMF27) — text, required')
field('Bid Rank (SMF28) — number')
field('Bid Validity Start (SMF29) — date')
field('Bid Validity End (SMF30) — date')
no_docs(); sep()

add_task(6, 'Approval For L1 Deviation', 'SMT11', True)
fields_label()
field('Bid Details (SMF27) — text, required')
field('Reason for Deviation — dropdown, required, NEW')
field('Remark — text, NEW')
no_docs(); sep()

doc.add_heading('ORIGIN', level=3)

add_task(7, 'Select Sailing Schedule', 'SMT12', False)
fields_label()
field('Sailing Schedule Details (SMF31) — text, required')
field('Sailing Date as per Schedule (SMF32) — date')
field('Vessel Name as per Schedule (SMF33) — text')
no_docs(); sep()

add_task(8, 'Approve Booking Note', 'SMT16', True)
fields_label()
field('Booking Number (SMF34) — text, required')
field('Sailing Date as per BN (SMF35) — date, required')
field('Vessel Name as per BN (SMF36) — text, required')
field('VGM Cut-off (SMF37) — datetime, required')
field('Gate Open (SMF38) — datetime, required')
field('Document Cut-off (SMF39) — datetime, required')
field('Bill Validity (SMF40) — date, required')
field('Gate-in Cut-off (SMF41) — datetime, required')
field('SI Cut-off (SMF42) — datetime, required')
field('Yard Details (SMF43) — text, required')
no_docs(); sep()

add_task(9, 'Approve Draft BL', 'SMT5', True)
fields_label()
field('Auto — Negotiable soft approval, no hardcoded fields')
no_docs(); sep()

doc.add_heading('IN TRANSIT', level=3)

add_task(10, 'Approve Final BL', 'SMT29', True)
fields_label()
field('Auto — Soft-approval task, no hardcoded fields')
no_docs(); sep()

add_task(11, 'Courier Docket Details Upload', 'SMT23', False)
fields_label()
field('AWB Number (SMF81) — text, required')
field('Courier Company (SMF82) — text, required')
field('Courier Receiving Address (SMF83) — text, required')
no_docs(); sep()

doc.add_heading('DESTINATION', level=3)

add_task(12, 'Upload Bill of Entry Details', 'SMT24', False)
fields_label()
field('BOE Number (SMF84) — text, required')
field('BOE Entry Port (SMF85) — text, required')
field('BOE Entry Date (SMF86) — date, required')
no_docs(); sep()

add_task(13, 'Vehicle Loading Confirmation', 'SMT3', False)
fields_label()
field('Container Loaded At — text (Negotiable document task)')
no_docs(); sep()

add_task(14, 'Detention Free Time', 'TBD', False, True)
fields_label()
field('Detention Free Time — text, required')
no_docs(); sep()

add_task(15, 'Confirm CFS Vendor', 'TBD', False, True)
fields_label()
field('CFS Vendors — dropdown, required')
no_docs(); sep()

add_task(16, 'Confirm ICD Vendor', 'TBD', False, True)
fields_label()
field('ICD Vendors — dropdown, required')
no_docs(); sep()

add_task(17, 'Transporter Confirmation', 'TBD', False, True)
fields_label()
field('Transporter — dropdown, required')
no_docs(); sep()

# ═══ FF ═══
doc.add_heading('FREIGHT FORWARDER (FF) — 10 Tasks', level=2)
doc.add_heading('ORIGIN', level=3)

add_task(18, 'Confirm Sailing Schedule', 'SMT13', False)
fields_label()
field('Sailing Schedule Details (SMF31) — text, required')
field('Sailing Date as per Schedule (SMF32) — date')
field('Vessel Name as per Schedule (SMF33) — text')
no_docs(); sep()

add_task(19, 'Upload Booking Note FCL', 'SMT32', False)
fields_label()
field('Booking Date Info same as Sailing Schedule (SMF96) — text')
doc_header('Booking Note (FCL)')
field('Booking Number (SMF34) — text, required')
field('Sailing Date (SMF112) — date')
field('Number of Containers (SMF111) — number')
field('Bill Validity (SMF40) — date')
field('SI Cut-off (SMF42) — datetime')
field('Document Cut-off (SMF39) — datetime')
field('VGM Cut-off (SMF37) — datetime')
field('Gate Open (SMF38) — datetime')
field('Gate-in Cut-off (SMF41) — datetime')
field('Vessel Name (SMF110) — text')
field('Yard Details (SMF43) — text')
field('Arrival Date (SMF89) — date')
sep()

add_task(20, 'Empty Container Pickup Details', 'SMT2', False)
fields_label()
field('Empty Container Pickup Date — date, required (Negotiable)')
no_docs(); sep()

add_task(21, 'Enter Container Weight Details', 'SMT20', False)
fields_label()
field('Gross Weight (SMF62) — number, required')
field('Net Weight (SMF63) — number, required')
field('Tare Weight (SMF64) — number, required')
field('UOM (SMF65) — dropdown, required')
field('Seal Number (SMF66) — text, required')
field('Container Type as per Tracking (SMF108) — text, required')
field('Container Size as per Tracking (SMF109) — text, required')
no_docs(); sep()

add_task(22, 'Upload Commercial Invoice & Packing List', 'SMT3', False)
fields_label()
field('Commercial Invoice — upload, required')
field('Packing List — upload, required')
no_docs(); sep()

add_task(23, 'Upload Draft BL', 'SMT3', True)
fields_label()
field('Draft BL — upload, required')
no_docs(); sep()

doc.add_heading('IN TRANSIT', level=3)

add_task(24, 'Upload Final BL & Freight Certificate', 'SMT19', True)
doc.add_paragraph('Fields: None')
doc_header('Final BL (FCL)')
field('House BL Number (SMF55) — text, required')
field('House BL Date (SMF56) — date, required')
field('Master BL Number (SMF45) — text, required')
field('Master BL Date (SMF46) — date, required')
field('Container Number List (SMF61) — text, required')
field('Actual Departure Date (SMF93) — date, required')
field('Vessel Name as per BL (SMF53) — text, required')
field('Vessel Number (SMF58) — text, required')
field('Voyage No (SMF54) — text, required')
field('Net Weight per BL (SMF51) — number, required')
field('Net Weight per BL UOM (SMF52) — dropdown, required')
field('Gross Weight per BL (SMF49) — number, required')
field('Gross Weight per BL UOM (SMF50) — dropdown, required')
field('Destination Arrival Date (SMF135) — date, required')
sep()

add_task(25, 'Upload Cargo Arrival Notice', 'SMT22', False)
doc.add_paragraph('Fields: None')
doc_header('CAN')
field('IGM Number (SMF74) — text, required')
field('IGM Item Number (SMF75) — text, required')
field('IGM Sub-item Number (SMF76) — text, required')
field('Inward Entry Date (SMF79) — date, required')
field('Gateway IGM Date (SMF80) — date, required')
field('Vessel IMO (SMF78) — text, required')
field('Vessel Code (SMF77) — text, required')
sep()

doc.add_heading('DESTINATION', level=3)

add_task(26, 'Upload Delivery Order', 'SMT21', False)
doc.add_paragraph('Fields: None')
doc_header('DO')
field('DO Number (SMF67) — text, required')
field('DO Date (SMF68) — date, required')
field('BL Number List (SMF73) — text, required')
field('DO Expiry Time (SMF70) — datetime, required')
field('Detention Free Expiry Time (SMF69) — datetime, required')
field('Laden Container Yard (SMF72) — text, required')
field('Empty Container Yard (SMF71) — text, required')
sep()

add_task(27, 'FF Incidental Events', 'TBD', False, True)
fields_label()
field('Incidental Charges — addmore, required')
field('Type of Charge — auto (Incidental / Self-Reimb / Third-Party)')
no_docs(); sep()

# ═══ CHA ═══
doc.add_heading('CHA — 1 Task', level=2)
doc.add_heading('DESTINATION', level=3)

add_task(28, 'CHA Incidental Events', 'TBD', False, True)
fields_label()
field('Incidental Charges — addmore, required')
field('Type of Charge — auto (Incidental / Self-Reimb / Third-Party)')
no_docs(); sep()

# ═══ CFS ═══
doc.add_heading('CFS — 4 Tasks', level=2)
doc.add_heading('DESTINATION', level=3)

add_task(29, 'CFS Gate In Date & Time', 'TBD', False, True)
fields_label()
field('CFS Gate In Date & Time — datetime, required')
no_docs(); sep()

add_task(30, 'CFS Destuff Indicator Confirmation', 'TBD', False, True)
fields_label()
field('CFS Destuff Indicator — text, required')
no_docs(); sep()

add_task(31, 'CFS Gate Out Date & Time', 'TBD', False, True)
fields_label()
field('CFS Gate Out Date & Time — datetime, required')
no_docs(); sep()

add_task(32, 'CFS Incidental Events', 'TBD', False, True)
fields_label()
field('Incidental Charges — addmore, required')
field('Type of Charge — auto')
no_docs(); sep()

# ═══ ICD ═══
doc.add_heading('ICD — 4 Tasks', level=2)
doc.add_heading('DESTINATION', level=3)

add_task(33, 'ICD Gate In Date & Time', 'TBD', False, True)
fields_label()
field('ICD Gate In Date & Time — datetime, required')
no_docs(); sep()

add_task(34, 'ICD Destuff Indicator Confirmation', 'TBD', False, True)
fields_label()
field('ICD Destuff Indicator — text, required')
no_docs(); sep()

add_task(35, 'ICD Gate Out Date & Time', 'TBD', False, True)
fields_label()
field('ICD Gate Out Date & Time — datetime, required')
no_docs(); sep()

add_task(36, 'ICD Incidental Events', 'TBD', False, True)
fields_label()
field('Incidental Charges — addmore, required')
field('Type of Charge — auto')
no_docs(); sep()

# ═══ TRANSPORTER ═══
doc.add_heading('TRANSPORTER — 4 Tasks', level=2)
doc.add_heading('DESTINATION', level=3)

add_task(37, 'Carrier Confirmation', 'TBD', False, True)
fields_label()
field('Carrier Confirmation Status — dropdown, required — Options: Confirmed, Pending, Rejected')
field('Vehicle Number — text, required')
field('Driver Name — text, required')
field('Driver Mobile — text, required')
no_docs(); sep()

add_task(38, 'Consignment Note & Eway Bill', 'TBD', False, True)
fields_label()
field('Consignment Note — text, required')
field('Eway Bill Number — text, required')
no_docs(); sep()

add_task(39, 'Empty Container Return Details', 'TBD', False, True)
fields_label()
field('Empty Container Return Date — date, required')
no_docs(); sep()

add_task(40, 'Transporter Incidental Events', 'TBD', False, True)
fields_label()
field('Incidental Charges — addmore, required')
field('Type of Charge — auto')
no_docs()

doc.save('D:/ShipsyShipmentExecutionDemoDev/Shipsy_EXIM_Manage_Tasks_Complete_List.docx')
print('Done!')
