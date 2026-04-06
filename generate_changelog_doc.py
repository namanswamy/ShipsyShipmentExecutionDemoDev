from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Shipsy EXIM - Manage Tasks Demo: Changes Summary', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Date: 7 April 2026\nLive Demo: https://shipment-execution-demo.vercel.app')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# Helper functions
def add_section(num, title_text):
    h = doc.add_heading(f'{num}. {title_text}', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x6E, 0xC3)

def add_sub(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(12)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_body(text):
    doc.add_paragraph(text)

# ── Section 1 ──
add_section(1, 'Spot / Normal Shipment Indicators')

add_sub('What changed')
add_bullet('Shipment cards now visually distinguish between Spot and Normal RFQ modes.')
add_bullet(' shipments have an orange left border, warm background gradient, and an orange "SPOT" badge in the card header.', 'Spot')
add_bullet(' shipments have a blue left border, cool background gradient, and a blue "NORMAL" badge in the card header.', 'Normal')

add_sub('Pre-configured shipments')
add_bullet('4 new shipment cards added with Task 1 ("Select Mode of Shipment") already completed:')
add_bullet('2 Spot shipments (SEHR38490, SEHR38489)')
add_bullet('2 Normal shipments (SEHR38488, SEHR38486)')

add_sub('Dynamic behavior')
add_bullet('When any shipment\'s Task 1 is submitted with a Spot/Normal selection, the card border and badge update live.')

# ── Section 2 ──
add_section(2, 'Mode Dropdown Filtering')

add_sub('What changed')
add_bullet('Mode dropdown names cleaned up: "BULK (MR)" → "Bulk", "BREAK BULK (MB)" → "Break Bulk", "AIR" → "Air".')
add_bullet('The Mode dropdown in Task 1 now shows conditional options based on the shipment card\'s mode:')
add_bullet(' shipments → FCL, LCL, Bulk, Break Bulk', 'FCL / LCL / Break Bulk')
add_bullet(' shipments → Air only', 'AIR')

add_sub('Shipment card changes')
add_bullet('Removed standalone BULK shipment card (Bulk only exists as a dropdown option).')
add_bullet('BB shipment card now displays as "Break Bulk".')

# ── Section 3 ──
add_section(3, 'Port Details Dropdown Update')

add_sub('What changed')
add_body('All 4 fields in Task 2 ("Select Port Details") — Place of Receipt at Origin, Port of Loading, Port of Discharge, Destination Port — now share the same 20 port options:')
add_body('Kutno, Pipavav, ICD/CFS Chennai, Nhava Sheva, Songkhla, Salvador BA, Santos SP, Vizagapatnam, Lat Krabang, Penang Port, Dadri (Greater Noida), Fremantle, Chittagong, Legnickie Pole Poland, Manaus AM, Allcargo Global Logistics Limited (CFS), Copenhagen, Rio de Janeiro RJ, Paranagua PR, Hong Kong.')

# ── Section 4 ──
add_section(4, 'Container Details Repeater')

add_sub('What changed')
add_bullet('Task 3 ("Enter Container Details"): Removed the "Container Details" text field.')
add_bullet('The remaining 5 fields (Container Type, Size, Count, Total Weight, Weight Unit) now display as a clean horizontal table row with numbered rows, column headers, compact inline inputs, and per-row remove button.')
add_bullet('"+ Add More" button adds additional container entry rows.')

# ── Section 5 ──
add_section(5, 'GPO Bid Card Color Theming')

add_sub('GPO Task (Run Global Plan Optimizer)')
add_bullet(' rate cards → warm orange background and borders.', 'Spot')
add_bullet(' rate cards → grey background and borders.', 'Normal')
add_bullet('Selection indicator bar uses orange for Spot, blue for Normal.')

add_sub('L1 Deviation Approval Task')
add_bullet('Same Spot/Normal color scheme but with lighter/muted tones to indicate read-only state.')

# ── Section 6 ──
add_section(6, 'Incidental Charges: BL / Container Level Mapping')

add_sub('What changed')
add_body('Each incidental charge now has a defined level — either BL level or Container level. Based on this, only the relevant table is shown (not both).')

# Table
table = doc.add_table(rows=8, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Grid Accent 1'

headers = ['Charge', 'Type', 'Level']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ('Loading charges', 'Incidental', 'BL (1 row)'),
    ('Storage charges', 'Incidental', 'Container (3 rows)'),
    ('Documentation charges', 'Self-Reimbursement', 'BL (1 row)'),
    ('Special equipment charges', 'Self-Reimbursement', 'Container (3 rows)'),
    ('Agency charges', 'Third-Party Reimbursement', 'BL (1 row)'),
    ('License charges', 'Third-Party Reimbursement', 'Container (3 rows)'),
    ('Registration charges', 'Third-Party Reimbursement', 'Container (3 rows)'),
]
for row_idx, (c, t, l) in enumerate(data):
    table.rows[row_idx + 1].cells[0].text = c
    table.rows[row_idx + 1].cells[1].text = t
    table.rows[row_idx + 1].cells[2].text = l

doc.add_paragraph()
add_bullet(' charges always show exactly 1 BL row (1 ASN = 1 BL).', 'BL-level')
add_bullet(' charges show only the Container table.', 'Container-level')

# ── Section 7 ──
add_section(7, 'GST Manual Input with Mutual Exclusion')

add_sub('What changed (Third-Party Reimbursement charges)')
add_bullet('CGST, SGST, and IGST are now manual editable input fields (previously auto-calculated).')
add_bullet(' auto-fills SGST with the same value when entered.', 'CGST')
add_bullet('Mutual exclusion logic:')
add_bullet('If CGST or SGST has a value → IGST field is disabled (greyed out).')
add_bullet('If IGST has a value → CGST and SGST fields are disabled (greyed out).')
add_bullet('Invoice Value auto-recalculates as: Basic Value + CGST + SGST + IGST.')

# ── Section 8 ──
add_section(8, 'Grand Total Breakdown by Charge Category')

add_sub('What changed')
add_body('The Grand Total section at the bottom of the Incidental Charges detail view now shows a 4-box row:')

t2 = doc.add_table(rows=1, cols=4)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2.style = 'Light Grid Accent 1'
labels = ['Incidental Total', 'Self-Reimbursement Total', 'Third-Party Reimbursement Total', 'Grand Total']
for i, l in enumerate(labels):
    cell = t2.rows[0].cells[i]
    cell.text = l
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)

doc.add_paragraph()
add_bullet('Each category total updates live as values are filled in.')
add_bullet('Category boxes are color-coded (orange, blue, green). Grand Total uses dark background.')

# ── Section 9 ──
add_section(9, 'Demo Incidental Shipments: Approver Response Flow')

add_body('4 new demo shipment cards demonstrate the full incidental charges lifecycle:')

add_sub('Demo Test Shipment Incidental Task 1 — Ready to Fill')
add_bullet('Incidental task ready to be filled by vendor. Standard flow.')

add_sub('Demo Test Shipment Incidental Task 2 — "Rework Required"')
add_bullet('Approver sent back with mixed actions:')
add_bullet(' → Approved (green badge, row locked/read-only)', 'Loading charges')
add_bullet(' → Rework (yellow badge + remark: "Please attach supporting invoice document")', 'Documentation charges')
add_bullet(' → Rejected (red badge + remark: "Rate not as per contract terms")', 'Agency charges')
add_bullet('Task status shows "Rework Required". Rework/Rejected rows remain editable.')

add_sub('Demo Test Shipment Incidental Task 3 — "Not Approved"')
add_bullet('All charges Rejected by approver with individual remarks.')
add_bullet('Task status shows "Not Approved". All fields editable for re-submission.')

add_sub('Demo Test Shipment Incidental Task 4 — "Approved"')
add_bullet('All charges Approved by approver with remarks.')
add_bullet('Task status shows "Approved". All fields locked/read-only.')

add_sub('UI details')
add_bullet(' column appears automatically when charge rows have approver data.', 'Approver Response')
add_bullet('Charge selection screen is locked for reviewed shipments — remove buttons hidden, dropdowns disabled.')
add_bullet('"Show Draft" button renamed to "View Approver Response" for shipments with approver feedback.')

# ── Section 10 ──
add_section(10, 'UI Fixes')

add_bullet('Status dropdown widened so "Rework Required" and other long names display fully. Applied uniformly to all task rows.')
add_bullet('Grid column alignment fixed — status column no longer overlaps the assignee column.')
add_bullet('Persona visibility bug fixed — switching between demo shipments no longer causes persona tabs to disappear.')

# Save
output_path = '/Users/shipsy/Desktop/RFQ Shipment Execution Demo/ShipsyShipmentExecutionDemoDev/Changes_Summary.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
