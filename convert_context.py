from docx import Document
from docx.shared import Pt

doc = Document()
with open('PROJECT_CONTEXT.md', 'r', encoding='utf-8') as f:
    in_code = False
    for line in f:
        line = line.rstrip()
        if line.startswith('```'):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(9)
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)

doc.save('PROJECT_CONTEXT.docx')
print('Done - saved as PROJECT_CONTEXT.docx')
